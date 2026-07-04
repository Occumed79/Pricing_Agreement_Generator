from __future__ import annotations

import hashlib
import io
import os
import re
import uuid
import zipfile
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Tuple

import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

try:
    import psycopg
    from psycopg.rows import dict_row
except Exception:  # pragma: no cover - app can still run with bundled templates locally
    psycopg = None
    dict_row = None

APP_TITLE = "Pricing Agreement Generator"
TEMPLATE_DIR = Path(__file__).parent / "templates"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
AUTO_DETECT_LABEL = "Auto-detect from provider address"


@dataclass
class TemplateMeta:
    id: str
    name: str
    category: str
    requires_code: bool
    description: str
    source: str  # "database", "default", or "session"
    filename: str
    is_default: bool = False
    path: Optional[Path] = None
    bytes_data: Optional[bytes] = None
    file_sha256: str = ""


DEFAULT_TEMPLATES: List[TemplateMeta] = [
    TemplateMeta(
        id="default-overseas-medical-pricing-agreement",
        name="Overseas Medical Pricing Agreement",
        category="Overseas",
        requires_code=True,
        description="International medical pricing agreement. Gray placeholders are filled with the ISO 4217 currency code selected or detected from the provider country.",
        source="default",
        filename="Overseas Medical Pricing Agreement.docx",
        is_default=True,
        path=TEMPLATE_DIR / "Overseas Medical Pricing Agreement.docx",
    ),
    TemplateMeta(
        id="default-medical-pricing-agreement-2026",
        name="1 - Medical Pricing Agreement - 2026",
        category="United States - Medical",
        requires_code=False,
        description="US medical pricing agreement.",
        source="default",
        filename="1 - Medical Pricing Agreement - 2026.docx",
        is_default=True,
        path=TEMPLATE_DIR / "1 - Medical Pricing Agreement - 2026.docx",
    ),
    TemplateMeta(
        id="default-dental-pricing-agreement-2025",
        name="2 - Dental Pricing Agreement -2025",
        category="United States - Dental",
        requires_code=False,
        description="US dental pricing agreement.",
        source="default",
        filename="2 - Dental Pricing Agreement -2025.docx",
        is_default=True,
        path=TEMPLATE_DIR / "2 - Dental Pricing Agreement -2025.docx",
    ),
    TemplateMeta(
        id="default-cardiovascular-components-2026",
        name="13 - Cardiovascular Components - 2026",
        category="United States - Cardiology",
        requires_code=False,
        description="US cardiology/cardiovascular pricing agreement.",
        source="default",
        filename="13 - Cardiovascular Components - 2026.docx",
        is_default=True,
        path=TEMPLATE_DIR / "13 - Cardiovascular Components - 2026.docx",
    ),
]

HIGHLIGHT_TO_FIELD = {
    WD_COLOR_INDEX.YELLOW: "name",
    WD_COLOR_INDEX.BRIGHT_GREEN: "address",
    WD_COLOR_INDEX.GREEN: "address",
    WD_COLOR_INDEX.TURQUOISE: "date",
    WD_COLOR_INDEX.BLUE: "date",
    WD_COLOR_INDEX.TEAL: "date",
    WD_COLOR_INDEX.GRAY_25: "code",
    WD_COLOR_INDEX.GRAY_50: "code",
}

FIELD_LABELS = {
    "name": "Yellow / Provider name",
    "address": "Green / Address",
    "date": "Blue or cyan / Date",
    "code": "Gray / Currency code",
}

# Offline country -> ISO 4217 currency code map. This avoids internet/API dependency.
# Countries sharing a currency intentionally point to the same code, e.g. Germany/France -> EUR.
COUNTRY_CURRENCY: Dict[str, str] = {
    "Afghanistan": "AFN",
    "Albania": "ALL",
    "Algeria": "DZD",
    "Andorra": "EUR",
    "Angola": "AOA",
    "Antigua and Barbuda": "XCD",
    "Argentina": "ARS",
    "Armenia": "AMD",
    "Australia": "AUD",
    "Austria": "EUR",
    "Azerbaijan": "AZN",
    "Bahamas": "BSD",
    "Bahrain": "BHD",
    "Bangladesh": "BDT",
    "Barbados": "BBD",
    "Belarus": "BYN",
    "Belgium": "EUR",
    "Belize": "BZD",
    "Benin": "XOF",
    "Bhutan": "BTN",
    "Bolivia": "BOB",
    "Bosnia and Herzegovina": "BAM",
    "Botswana": "BWP",
    "Brazil": "BRL",
    "Brunei": "BND",
    "Bulgaria": "BGN",
    "Burkina Faso": "XOF",
    "Burundi": "BIF",
    "Cabo Verde": "CVE",
    "Cambodia": "KHR",
    "Cameroon": "XAF",
    "Canada": "CAD",
    "Central African Republic": "XAF",
    "Chad": "XAF",
    "Chile": "CLP",
    "China": "CNY",
    "Colombia": "COP",
    "Comoros": "KMF",
    "Congo": "XAF",
    "Costa Rica": "CRC",
    "Croatia": "EUR",
    "Cuba": "CUP",
    "Cyprus": "EUR",
    "Czech Republic": "CZK",
    "Denmark": "DKK",
    "Djibouti": "DJF",
    "Dominica": "XCD",
    "Dominican Republic": "DOP",
    "Ecuador": "USD",
    "Egypt": "EGP",
    "El Salvador": "USD",
    "Equatorial Guinea": "XAF",
    "Eritrea": "ERN",
    "Estonia": "EUR",
    "Eswatini": "SZL",
    "Ethiopia": "ETB",
    "Fiji": "FJD",
    "Finland": "EUR",
    "France": "EUR",
    "Gabon": "XAF",
    "Gambia": "GMD",
    "Georgia": "GEL",
    "Germany": "EUR",
    "Ghana": "GHS",
    "Greece": "EUR",
    "Grenada": "XCD",
    "Guatemala": "GTQ",
    "Guinea": "GNF",
    "Guinea-Bissau": "XOF",
    "Guyana": "GYD",
    "Haiti": "HTG",
    "Honduras": "HNL",
    "Hong Kong": "HKD",
    "Hungary": "HUF",
    "Iceland": "ISK",
    "India": "INR",
    "Indonesia": "IDR",
    "Iran": "IRR",
    "Iraq": "IQD",
    "Ireland": "EUR",
    "Israel": "ILS",
    "Italy": "EUR",
    "Ivory Coast": "XOF",
    "Jamaica": "JMD",
    "Japan": "JPY",
    "Jordan": "JOD",
    "Kazakhstan": "KZT",
    "Kenya": "KES",
    "Kuwait": "KWD",
    "Kyrgyzstan": "KGS",
    "Laos": "LAK",
    "Latvia": "EUR",
    "Lebanon": "LBP",
    "Lesotho": "LSL",
    "Liberia": "LRD",
    "Libya": "LYD",
    "Liechtenstein": "CHF",
    "Lithuania": "EUR",
    "Luxembourg": "EUR",
    "Macau": "MOP",
    "Madagascar": "MGA",
    "Malawi": "MWK",
    "Malaysia": "MYR",
    "Maldives": "MVR",
    "Mali": "XOF",
    "Malta": "EUR",
    "Mauritania": "MRU",
    "Mauritius": "MUR",
    "Mexico": "MXN",
    "Moldova": "MDL",
    "Monaco": "EUR",
    "Mongolia": "MNT",
    "Montenegro": "EUR",
    "Morocco": "MAD",
    "Mozambique": "MZN",
    "Myanmar": "MMK",
    "Namibia": "NAD",
    "Nepal": "NPR",
    "Netherlands": "EUR",
    "New Zealand": "NZD",
    "Nicaragua": "NIO",
    "Niger": "XOF",
    "Nigeria": "NGN",
    "North Macedonia": "MKD",
    "Norway": "NOK",
    "Oman": "OMR",
    "Pakistan": "PKR",
    "Panama": "PAB",
    "Papua New Guinea": "PGK",
    "Paraguay": "PYG",
    "Peru": "PEN",
    "Philippines": "PHP",
    "Poland": "PLN",
    "Portugal": "EUR",
    "Qatar": "QAR",
    "Romania": "RON",
    "Russia": "RUB",
    "Rwanda": "RWF",
    "Saint Kitts and Nevis": "XCD",
    "Saint Lucia": "XCD",
    "Saint Vincent and the Grenadines": "XCD",
    "Samoa": "WST",
    "San Marino": "EUR",
    "Saudi Arabia": "SAR",
    "Senegal": "XOF",
    "Serbia": "RSD",
    "Seychelles": "SCR",
    "Sierra Leone": "SLE",
    "Singapore": "SGD",
    "Slovakia": "EUR",
    "Slovenia": "EUR",
    "Solomon Islands": "SBD",
    "Somalia": "SOS",
    "South Africa": "ZAR",
    "South Korea": "KRW",
    "Spain": "EUR",
    "Sri Lanka": "LKR",
    "Sudan": "SDG",
    "Suriname": "SRD",
    "Sweden": "SEK",
    "Switzerland": "CHF",
    "Syria": "SYP",
    "Taiwan": "TWD",
    "Tajikistan": "TJS",
    "Tanzania": "TZS",
    "Thailand": "THB",
    "Togo": "XOF",
    "Tonga": "TOP",
    "Trinidad and Tobago": "TTD",
    "Tunisia": "TND",
    "Turkey": "TRY",
    "Turkmenistan": "TMT",
    "Uganda": "UGX",
    "Ukraine": "UAH",
    "United Arab Emirates": "AED",
    "United Kingdom": "GBP",
    "United States": "USD",
    "Uruguay": "UYU",
    "Uzbekistan": "UZS",
    "Vanuatu": "VUV",
    "Venezuela": "VES",
    "Vietnam": "VND",
    "Yemen": "YER",
    "Zambia": "ZMW",
    "Zimbabwe": "ZWL",
}

COUNTRY_ALIASES: Dict[str, str] = {
    "uae": "United Arab Emirates",
    "u.a.e.": "United Arab Emirates",
    "emirates": "United Arab Emirates",
    "uk": "United Kingdom",
    "u.k.": "United Kingdom",
    "great britain": "United Kingdom",
    "britain": "United Kingdom",
    "england": "United Kingdom",
    "scotland": "United Kingdom",
    "wales": "United Kingdom",
    "northern ireland": "United Kingdom",
    "usa": "United States",
    "u.s.a.": "United States",
    "us": "United States",
    "u.s.": "United States",
    "america": "United States",
    "united states of america": "United States",
    "new zealand": "New Zealand",
    "nz": "New Zealand",
    "aotearoa": "New Zealand",
    "korea": "South Korea",
    "republic of korea": "South Korea",
    "south korea": "South Korea",
    "viet nam": "Vietnam",
    "cote d'ivoire": "Ivory Coast",
    "côte d’ivoire": "Ivory Coast",
    "côte d'ivoire": "Ivory Coast",
    "ivory coast": "Ivory Coast",
    "czechia": "Czech Republic",
    "macedonia": "North Macedonia",
    "russian federation": "Russia",
    "hong kong sar": "Hong Kong",
    "hong kong, china": "Hong Kong",
    "macau sar": "Macau",
    "macao": "Macau",
}

# Longest first prevents matching "Congo" inside a longer country if both exist.
COUNTRY_PATTERNS: List[Tuple[str, str]] = sorted(
    [(name.lower(), name) for name in COUNTRY_CURRENCY] + [(alias, canonical) for alias, canonical in COUNTRY_ALIASES.items()],
    key=lambda item: len(item[0]),
    reverse=True,
)


# -----------------------------
# UI styling
# -----------------------------
def inject_css() -> None:
    st.markdown(
        """
        <style>
        :root {
            --glass-bg: rgba(13, 24, 43, 0.54);
            --glass-border: rgba(255, 255, 255, 0.22);
            --glow-a: rgba(125, 211, 252, 0.42);
            --glow-b: rgba(168, 85, 247, 0.34);
            --glow-c: rgba(34, 211, 238, 0.26);
            --text-muted: rgba(226, 232, 240, 0.80);
        }

        .stApp {
            background:
                radial-gradient(circle at 10% 15%, rgba(125, 211, 252, 0.30), transparent 31%),
                radial-gradient(circle at 86% 10%, rgba(168, 85, 247, 0.27), transparent 34%),
                radial-gradient(circle at 58% 84%, rgba(20, 184, 166, 0.21), transparent 29%),
                linear-gradient(135deg, #06101f 0%, #0a1020 52%, #0b1325 100%);
            color: #f8fafc;
        }

        .block-container {
            padding-top: 2.4rem;
            padding-bottom: 4rem;
            max-width: 1180px;
        }

        [data-testid="stHeader"] { background: transparent; }

        .hero-card, .glass-card, .metric-card, .status-pill {
            position: relative;
            border: 1px solid var(--glass-border);
            background: linear-gradient(135deg, rgba(255,255,255,0.15), rgba(255,255,255,0.06));
            box-shadow:
                0 0 0 1px rgba(255,255,255,0.06) inset,
                0 18px 55px rgba(0,0,0,0.30),
                0 0 34px var(--glow-a),
                0 0 76px var(--glow-b),
                0 0 110px var(--glow-c);
            backdrop-filter: blur(23px) saturate(185%);
            -webkit-backdrop-filter: blur(23px) saturate(185%);
            border-radius: 28px;
        }

        .hero-card {
            padding: 30px 34px;
            margin-bottom: 22px;
            overflow: hidden;
        }

        .hero-card:before {
            content: "";
            position: absolute;
            inset: -2px;
            background:
                linear-gradient(90deg, transparent, rgba(255,255,255,0.38), transparent),
                radial-gradient(circle at 20% 0%, rgba(125,211,252,0.42), transparent 34%);
            opacity: 0.48;
            pointer-events: none;
        }

        .hero-title {
            font-size: clamp(2.1rem, 5vw, 4.3rem);
            line-height: 0.95;
            font-weight: 850;
            letter-spacing: -0.06em;
            margin: 0 0 12px;
            color: #ffffff;
            text-shadow: 0 0 38px rgba(125, 211, 252, 0.68);
        }

        .hero-subtitle {
            max-width: 830px;
            color: var(--text-muted);
            font-size: 1.05rem;
            margin: 0;
        }

        .glass-card {
            padding: 18px 20px;
            margin: 14px 0 18px;
        }

        .glass-card h3 { margin-top: 0; margin-bottom: 8px; }
        .glass-card p, .glass-card li { color: var(--text-muted); }

        .status-pill {
            display: inline-flex;
            align-items: center;
            gap: 8px;
            padding: 9px 13px;
            border-radius: 999px;
            font-weight: 750;
            margin: 2px 0 12px;
        }

        .metric-grid {
            display: grid;
            grid-template-columns: repeat(4, minmax(0, 1fr));
            gap: 12px;
            margin: 12px 0 18px;
        }

        .metric-card {
            padding: 14px 16px;
            border-radius: 20px;
        }

        .metric-value {
            font-size: 1.9rem;
            font-weight: 800;
            letter-spacing: -0.04em;
        }

        .metric-label { color: var(--text-muted); font-size: 0.86rem; }

        div[data-testid="stTabs"] button { border-radius: 999px !important; }

        .stButton > button, .stDownloadButton > button {
            border-radius: 999px;
            border: 1px solid rgba(255,255,255,0.24);
            background: linear-gradient(135deg, rgba(56,189,248,0.92), rgba(168,85,247,0.76));
            color: white;
            box-shadow: 0 0 26px rgba(125,211,252,0.42), 0 18px 36px rgba(0,0,0,0.28);
            font-weight: 760;
        }

        .stButton > button:hover, .stDownloadButton > button:hover {
            border-color: rgba(255,255,255,0.58);
            box-shadow: 0 0 38px rgba(125,211,252,0.60), 0 22px 44px rgba(0,0,0,0.34);
        }

        [data-testid="stDataFrame"] {
            border-radius: 20px;
            overflow: hidden;
            border: 1px solid rgba(255,255,255,0.16);
        }

        code {
            border-radius: 8px;
            padding: 0.15rem 0.35rem;
            background: rgba(255,255,255,0.12) !important;
        }

        @media (max-width: 760px) {
            .metric-grid { grid-template-columns: repeat(2, minmax(0, 1fr)); }
            .hero-card { padding: 24px 22px; }
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


def glass_card(title: str, body_html: str) -> None:
    st.markdown(
        f"""
        <div class="glass-card">
            <h3>{title}</h3>
            {body_html}
        </div>
        """,
        unsafe_allow_html=True,
    )


def status_pill(text: str) -> None:
    st.markdown(f"<div class='status-pill'>{text}</div>", unsafe_allow_html=True)


def metrics_card(counts: Dict[str, int]) -> None:
    st.markdown(
        f"""
        <div class="metric-grid">
          <div class="metric-card"><div class="metric-value">{counts.get('name', 0)}</div><div class="metric-label">Yellow name markers</div></div>
          <div class="metric-card"><div class="metric-value">{counts.get('address', 0)}</div><div class="metric-label">Green address markers</div></div>
          <div class="metric-card"><div class="metric-value">{counts.get('date', 0)}</div><div class="metric-label">Blue date markers</div></div>
          <div class="metric-card"><div class="metric-value">{counts.get('code', 0)}</div><div class="metric-label">Gray currency markers</div></div>
        </div>
        """,
        unsafe_allow_html=True,
    )


# -----------------------------
# General helpers
# -----------------------------
def today_label() -> str:
    today = date.today()
    return f"{today.strftime('%B')} {today.day}, {today.year}"


def safe_filename(value: str, max_length: int = 120) -> str:
    value = re.sub(r"[\\/:*?\"<>|]+", "-", value)
    value = re.sub(r"\s+", " ", value).strip()
    value = value.strip(". ")
    return (value[:max_length].rstrip() or "Agreement")


def slugify(value: str) -> str:
    value = re.sub(r"[^a-zA-Z0-9]+", "-", value.lower()).strip("-")
    return value or "template"


def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def normalize_country(value: str) -> Optional[str]:
    if not value or str(value).strip().lower() in {"nan", "none", "null"}:
        return None
    raw = re.sub(r"\s+", " ", str(value).strip())
    lowered = raw.lower().strip(".,;:()[]{}")
    if lowered in COUNTRY_ALIASES:
        return COUNTRY_ALIASES[lowered]
    for country in COUNTRY_CURRENCY:
        if lowered == country.lower():
            return country
    return None


def detect_country_from_text(text: str) -> Optional[str]:
    if not text:
        return None
    haystack = f" {re.sub(r'[^A-Za-zÀ-ÿ.\'’ -]+', ' ', text).lower()} "
    haystack = re.sub(r"\s+", " ", haystack)
    for pattern, canonical in COUNTRY_PATTERNS:
        pattern_re = re.escape(pattern).replace(r"\ ", r"\s+")
        if re.search(rf"(?<![a-z]){pattern_re}(?![a-z])", haystack):
            return canonical
    return None


def resolve_currency_code(country: Optional[str]) -> str:
    if not country:
        return ""
    canonical = normalize_country(country) or country
    return COUNTRY_CURRENCY.get(canonical, "")


def prepare_records_for_template(records: pd.DataFrame, template: TemplateMeta, selected_country: str) -> pd.DataFrame:
    """Add country/currency preview columns for overseas templates."""
    out = records.copy()
    if not template.requires_code:
        return out

    resolved_countries: List[str] = []
    currency_codes: List[str] = []
    manual_country = selected_country if selected_country != AUTO_DETECT_LABEL else ""

    for _, row in out.iterrows():
        country = normalize_country(manual_country)
        if not country and "country" in out.columns:
            country = normalize_country(str(row.get("country", "")))
        if not country:
            country = detect_country_from_text(str(row.get("address", "")))
        code = resolve_currency_code(country)
        resolved_countries.append(country or "")
        currency_codes.append(code or "")

    out["resolved_country"] = resolved_countries
    out["currency_code"] = currency_codes
    return out


# -----------------------------
# Neon template library
# -----------------------------
def get_database_url() -> str:
    if "DATABASE_URL" in st.secrets:
        return str(st.secrets["DATABASE_URL"])
    return os.getenv("DATABASE_URL", "")


def database_configured() -> bool:
    return bool(get_database_url().strip()) and psycopg is not None


def get_db_connection():
    if not database_configured():
        raise RuntimeError("DATABASE_URL is not configured or psycopg is not installed.")
    return psycopg.connect(get_database_url(), row_factory=dict_row, autocommit=True)


def ensure_template_table() -> None:
    with get_db_connection() as conn:
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS agreement_templates (
                id TEXT PRIMARY KEY,
                name TEXT NOT NULL,
                category TEXT NOT NULL,
                requires_code BOOLEAN NOT NULL DEFAULT FALSE,
                description TEXT NOT NULL DEFAULT '',
                filename TEXT NOT NULL,
                content_type TEXT NOT NULL DEFAULT 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                file_bytes BYTEA NOT NULL,
                file_sha256 TEXT NOT NULL,
                is_default BOOLEAN NOT NULL DEFAULT FALSE,
                is_active BOOLEAN NOT NULL DEFAULT TRUE,
                created_at TIMESTAMPTZ NOT NULL DEFAULT NOW(),
                updated_at TIMESTAMPTZ NOT NULL DEFAULT NOW()
            );
            """
        )


def seed_default_templates_to_neon() -> None:
    """Keep the four bundled templates permanently available in Neon.

    This upserts by stable IDs so redeploys do not create duplicates.
    If the bundled template file changes, the Neon row is updated.
    """
    ensure_template_table()
    with get_db_connection() as conn:
        for template in DEFAULT_TEMPLATES:
            if not template.path or not template.path.exists():
                continue
            data = template.path.read_bytes()
            conn.execute(
                """
                INSERT INTO agreement_templates
                    (id, name, category, requires_code, description, filename, content_type, file_bytes, file_sha256, is_default, is_active)
                VALUES
                    (%(id)s, %(name)s, %(category)s, %(requires_code)s, %(description)s, %(filename)s, %(content_type)s, %(file_bytes)s, %(file_sha256)s, TRUE, TRUE)
                ON CONFLICT (id) DO UPDATE SET
                    name = EXCLUDED.name,
                    category = EXCLUDED.category,
                    requires_code = EXCLUDED.requires_code,
                    description = EXCLUDED.description,
                    filename = EXCLUDED.filename,
                    content_type = EXCLUDED.content_type,
                    file_bytes = EXCLUDED.file_bytes,
                    file_sha256 = EXCLUDED.file_sha256,
                    is_default = TRUE,
                    is_active = TRUE,
                    updated_at = NOW();
                """,
                {
                    "id": template.id,
                    "name": template.name,
                    "category": template.category,
                    "requires_code": template.requires_code,
                    "description": template.description,
                    "filename": template.filename,
                    "content_type": DOCX_MIME,
                    "file_bytes": data,
                    "file_sha256": sha256_bytes(data),
                },
            )


def load_templates_from_neon() -> List[TemplateMeta]:
    seed_default_templates_to_neon()
    with get_db_connection() as conn:
        rows = conn.execute(
            """
            SELECT id, name, category, requires_code, description, filename, file_bytes, file_sha256, is_default
            FROM agreement_templates
            WHERE is_active = TRUE
            ORDER BY is_default DESC, category ASC, name ASC;
            """
        ).fetchall()

    templates: List[TemplateMeta] = []
    for row in rows:
        templates.append(
            TemplateMeta(
                id=row["id"],
                name=row["name"],
                category=row["category"],
                requires_code=bool(row["requires_code"]),
                description=row["description"] or "",
                source="database",
                filename=row["filename"],
                is_default=bool(row["is_default"]),
                bytes_data=bytes(row["file_bytes"]),
                file_sha256=row["file_sha256"] or "",
            )
        )
    return templates


def save_template_to_neon(
    *,
    name: str,
    category: str,
    requires_code: bool,
    description: str,
    filename: str,
    template_bytes: bytes,
) -> str:
    ensure_template_table()
    template_id = f"custom-{slugify(name)}-{uuid.uuid4().hex[:8]}"
    with get_db_connection() as conn:
        conn.execute(
            """
            INSERT INTO agreement_templates
                (id, name, category, requires_code, description, filename, content_type, file_bytes, file_sha256, is_default, is_active)
            VALUES
                (%(id)s, %(name)s, %(category)s, %(requires_code)s, %(description)s, %(filename)s, %(content_type)s, %(file_bytes)s, %(file_sha256)s, FALSE, TRUE);
            """,
            {
                "id": template_id,
                "name": name,
                "category": category,
                "requires_code": requires_code,
                "description": description,
                "filename": filename,
                "content_type": DOCX_MIME,
                "file_bytes": template_bytes,
                "file_sha256": sha256_bytes(template_bytes),
            },
        )
    return template_id


def get_default_templates_from_files() -> List[TemplateMeta]:
    templates: List[TemplateMeta] = []
    for template in DEFAULT_TEMPLATES:
        if template.path and template.path.exists():
            templates.append(template)
    return templates


def get_all_templates() -> Tuple[List[TemplateMeta], str]:
    """Return templates and a source status string."""
    if database_configured():
        try:
            templates = load_templates_from_neon()
            return templates, "Neon template library connected. Templates are saved permanently."
        except Exception as exc:
            st.session_state["db_error"] = str(exc)
            templates = get_default_templates_from_files() + st.session_state.get("custom_templates", [])
            return templates, "Neon is configured, but the app could not connect. Falling back to bundled/session templates."

    templates = get_default_templates_from_files() + st.session_state.get("custom_templates", [])
    return templates, "Neon is not configured. Bundled templates work, but uploaded custom templates are session-only."


def get_template_bytes(template: TemplateMeta) -> bytes:
    if template.bytes_data is not None:
        return template.bytes_data
    if template.path:
        return template.path.read_bytes()
    raise ValueError("Template bytes are unavailable.")


# -----------------------------
# DOCX processing helpers
# -----------------------------
def classify_highlight(highlight) -> Optional[str]:
    return HIGHLIGHT_TO_FIELD.get(highlight)


def iter_table_paragraphs(table) -> Iterable:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested_table in cell.tables:
                yield from iter_table_paragraphs(nested_table)


def iter_document_paragraphs(doc: Document) -> Iterable:
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        yield from iter_table_paragraphs(table)
    for section in doc.sections:
        for part in (section.header, section.footer):
            for paragraph in part.paragraphs:
                yield paragraph
            for table in part.tables:
                yield from iter_table_paragraphs(table)


def scan_template(template_bytes: bytes) -> Dict[str, int]:
    doc = Document(io.BytesIO(template_bytes))
    counts = {"name": 0, "address": 0, "date": 0, "code": 0}
    for paragraph in iter_document_paragraphs(doc):
        for run in paragraph.runs:
            if not run.text.strip():
                continue
            field = classify_highlight(run.font.highlight_color)
            if field:
                counts[field] += 1
    return counts


def collect_highlighted_runs(doc: Document) -> Dict[str, List]:
    highlighted_runs = {"name": [], "address": [], "date": [], "code": [], "address_indexed": []}
    for paragraph_index, paragraph in enumerate(iter_document_paragraphs(doc)):
        for run in paragraph.runs:
            field = classify_highlight(run.font.highlight_color)
            if field:
                highlighted_runs[field].append(run)
                if field == "address":
                    highlighted_runs["address_indexed"].append((paragraph_index, run))
    return highlighted_runs


def set_run_text_preserve_linebreaks(run, text: str) -> None:
    """Replace a run while preserving Word line breaks for pasted multi-line addresses."""
    run.text = ""
    parts = str(text).split("\n")
    for i, part in enumerate(parts):
        if i > 0:
            run.add_break()
        run.add_text(part)


def clear_highlight(run) -> None:
    run.font.highlight_color = None


def split_address_groups(address_indexed_runs: List[Tuple[int, object]]) -> List[List]:
    """Group address placeholders that appear together in the same header/block.

    Some templates repeat the provider header on more than one page. Each header has its own
    Address 1 / Address 2 / Address 3 markers. This groups nearby green markers so the full
    address is filled once per header block instead of only once globally.
    """
    if not address_indexed_runs:
        return []

    groups: List[List] = []
    current: List = []
    last_index: Optional[int] = None
    for paragraph_index, run in address_indexed_runs:
        if last_index is None or paragraph_index - last_index <= 2:
            current.append(run)
        else:
            groups.append(current)
            current = [run]
        last_index = paragraph_index
    if current:
        groups.append(current)
    return groups


def replace_single_address_group(address_runs: List, address: str) -> None:
    """Fill one group of Address 1 / Address 2 / Address 3 style markers."""
    if not address_runs:
        return
    lines = [line.strip() for line in str(address).splitlines() if line.strip()]
    if not lines:
        lines = [str(address).strip()]

    if len(address_runs) == 1:
        assignments = ["\n".join(lines)]
    elif len(lines) <= len(address_runs):
        assignments = lines + [""] * (len(address_runs) - len(lines))
    else:
        assignments = lines[: len(address_runs) - 1]
        assignments.append("\n".join(lines[len(address_runs) - 1 :]))

    for run, value in zip(address_runs, assignments):
        set_run_text_preserve_linebreaks(run, value)
        clear_highlight(run)
    for run in address_runs[len(assignments) :]:
        set_run_text_preserve_linebreaks(run, "")
        clear_highlight(run)


def replace_address_markers(address_indexed_runs: List[Tuple[int, object]], address: str) -> None:
    for group in split_address_groups(address_indexed_runs):
        replace_single_address_group(group, address)


def replace_markers_in_docx(
    template_bytes: bytes,
    clinic_name: str,
    address: str,
    effective_date: str,
    currency_code: str = "",
) -> bytes:
    doc = Document(io.BytesIO(template_bytes))
    highlighted_runs = collect_highlighted_runs(doc)

    for field, value in {
        "name": clinic_name,
        "date": effective_date,
        "code": currency_code,
    }.items():
        for run in highlighted_runs.get(field, []):
            set_run_text_preserve_linebreaks(run, value)
            clear_highlight(run)

    replace_address_markers(highlighted_runs.get("address_indexed", []), address)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def build_zip(records: pd.DataFrame, template: TemplateMeta) -> bytes:
    template_bytes = get_template_bytes(template)
    generated_on = date.today().isoformat()
    effective_date = today_label()

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
        for _, row in records.iterrows():
            clinic_name = str(row["clinic_name"]).strip()
            address = str(row["address"]).strip()
            currency_code = str(row.get("currency_code", "")).strip().upper() if template.requires_code else ""
            docx_bytes = replace_markers_in_docx(
                template_bytes=template_bytes,
                clinic_name=clinic_name,
                address=address,
                effective_date=effective_date,
                currency_code=currency_code,
            )
            filename = safe_filename(f"{clinic_name} - {template.name} - {generated_on}.docx")
            zf.writestr(filename, docx_bytes)

    zip_buffer.seek(0)
    return zip_buffer.getvalue()


# -----------------------------
# Input parsing
# -----------------------------
def parse_spreadsheet(uploaded_file) -> Tuple[Optional[pd.DataFrame], Optional[str]]:
    if uploaded_file is None:
        return None, None

    try:
        name = uploaded_file.name.lower()
        if name.endswith(".csv"):
            df = pd.read_csv(uploaded_file)
        else:
            df = pd.read_excel(uploaded_file)
    except Exception as exc:
        return None, f"Could not read the spreadsheet: {exc}"

    df.columns = [str(col).strip() for col in df.columns]
    required = {"clinic_name", "address"}
    missing = sorted(required - set(df.columns))
    if missing:
        return None, "Missing required header(s): " + ", ".join(f"`{item}`" for item in missing)

    keep_cols = ["clinic_name", "address"]
    if "country" in df.columns:
        keep_cols.append("country")
    clean = df[keep_cols].copy()
    clean["clinic_name"] = clean["clinic_name"].astype(str).str.strip()
    clean["address"] = clean["address"].astype(str).str.strip()
    if "country" in clean.columns:
        clean["country"] = clean["country"].astype(str).str.strip()
    clean = clean[(clean["clinic_name"] != "") & (clean["address"] != "")]

    if clean.empty:
        return None, "The spreadsheet has the right headers, but no usable provider rows were found."

    return clean.reset_index(drop=True), None


def parse_pasted_providers(text: str) -> Tuple[Optional[pd.DataFrame], Optional[str]]:
    if not text.strip():
        return None, None

    blocks = re.split(r"\n\s*\n+", text.strip())
    rows = []
    skipped = []
    for index, block in enumerate(blocks, start=1):
        lines = [line.strip() for line in block.splitlines() if line.strip()]
        if len(lines) < 2:
            skipped.append(index)
            continue
        rows.append({"clinic_name": lines[0], "address": "\n".join(lines[1:])})

    if not rows:
        return None, "No valid provider blocks were found. Each provider needs a clinic name and at least one address line."

    df = pd.DataFrame(rows)
    if skipped:
        return df, f"Parsed {len(df)} provider(s). Skipped block(s) {', '.join(map(str, skipped))} because they did not include an address."
    return df, None


# -----------------------------
# App pages
# -----------------------------
def render_generate_tab() -> None:
    templates, status = get_all_templates()

    st.subheader("Generate Agreements")
    status_pill(status)
    if st.session_state.get("db_error"):
        with st.expander("Neon connection detail", expanded=False):
            st.code(st.session_state["db_error"])

    glass_card(
        "Quick flow",
        "<p>Select a saved template, add providers by spreadsheet or paste box, preview the list, then generate a downloadable ZIP. Templates are permanent in Neon; output files are temporary and are not saved.</p>",
    )

    if not templates:
        st.error("No templates were found. Add the default templates to the templates folder or connect Neon.")
        return

    selected_template = st.selectbox(
        "Select template",
        templates,
        format_func=lambda item: f"{item.category} — {item.name}",
    )

    with st.expander("Template details", expanded=False):
        st.write(selected_template.description or "No description provided.")
        st.write("Template source:", selected_template.source)
        st.write("Requires currency code:", "Yes" if selected_template.requires_code else "No")
        if selected_template.file_sha256:
            st.write("SHA-256:", selected_template.file_sha256)
        try:
            counts = scan_template(get_template_bytes(selected_template))
            metrics_card(counts)
        except Exception as exc:
            st.warning(f"Could not scan this template: {exc}")

    selected_country = AUTO_DETECT_LABEL
    if selected_template.requires_code:
        glass_card(
            "Overseas currency handling",
            "<p>The gray highlighted template markers are filled automatically with the ISO 4217 currency code. No one needs to type the three-letter code.</p><p>Select one country for the whole batch, or let the app detect the country from each provider address.</p>",
        )
        country_options = [AUTO_DETECT_LABEL] + sorted(COUNTRY_CURRENCY.keys())
        selected_country = st.selectbox("Country for currency code", country_options, index=0)
        if selected_country != AUTO_DETECT_LABEL:
            st.info(f"Resolved currency code: {COUNTRY_CURRENCY[selected_country]}")

    input_method = st.radio(
        "Provider input method",
        ["Upload Excel/CSV", "Paste providers manually"],
        horizontal=True,
    )

    records = None
    message = None

    if input_method == "Upload Excel/CSV":
        glass_card(
            "Spreadsheet requirements",
            "<p>Your file must include these exact headers:</p><ul><li><code>clinic_name</code></li><li><code>address</code></li></ul><p>Optional for overseas: <code>country</code>. Each row becomes one generated agreement.</p>",
        )
        upload = st.file_uploader("Upload Excel or CSV", type=["xlsx", "csv"], key="provider_sheet")
        records, message = parse_spreadsheet(upload)
    else:
        glass_card(
            "Paste format",
            "<p>First line of each block is <code>clinic_name</code>. Remaining lines are the full address. Use a blank line between providers.</p><p>For overseas templates, include the country in the address or select the country above.</p>",
        )
        sample = "Health Tick\nLevel 1, 19 Byron Avenue\nTakapuna, Auckland 0622\nNew Zealand\n\nABC Medical Clinic\n123 Main Street\nFresno, CA 93711"
        text = st.text_area("Paste providers", height=230, placeholder=sample)
        records, message = parse_pasted_providers(text)

    if message:
        if records is not None and not records.empty:
            st.info(message)
        else:
            st.warning(message)

    if records is not None and not records.empty:
        records_for_generation = prepare_records_for_template(records, selected_template, selected_country)

        st.markdown(f"### Preview — {len(records_for_generation)} agreement(s) will be generated")
        st.dataframe(records_for_generation, use_container_width=True, hide_index=True)

        can_generate = True
        if selected_template.requires_code:
            unresolved = records_for_generation[records_for_generation["currency_code"].astype(str).str.strip() == ""]
            if not unresolved.empty:
                can_generate = False
                st.error(
                    "The app could not resolve a currency code for every provider. Select one country for the whole batch or add the country to each address."
                )

        if st.button("Generate Agreements ZIP", type="primary", disabled=not can_generate):
            try:
                zip_bytes = build_zip(records_for_generation, selected_template)
                zip_name = f"Generated Agreements - {date.today().isoformat()}.zip"
                st.session_state["latest_zip"] = zip_bytes
                st.session_state["latest_zip_name"] = zip_name
                st.success("Agreements generated. Download the ZIP below.")
            except Exception as exc:
                st.error(f"Generation failed: {exc}")

    if st.session_state.get("latest_zip"):
        st.download_button(
            "Download generated ZIP",
            data=st.session_state["latest_zip"],
            file_name=st.session_state.get("latest_zip_name", "Generated Agreements.zip"),
            mime="application/zip",
        )


def render_configure_tab() -> None:
    st.subheader("Configure New Agreement Template")
    templates, status = get_all_templates()
    status_pill(status)

    glass_card(
        "Before uploading your template",
        """
        <ul>
          <li>File must be <code>.docx</code>. Do not upload <code>.doc</code>.</li>
          <li>Highlight provider/clinic name placeholders in <strong>yellow</strong>.</li>
          <li>Highlight address placeholders in <strong>green</strong>.</li>
          <li>Highlight date placeholders in <strong>blue/light blue</strong>.</li>
          <li>Highlight currency/code placeholders in <strong>gray</strong> only if the template needs them.</li>
          <li>Do not highlight labels. Only highlight the exact text that should be replaced.</li>
        </ul>
        <p>When Neon is connected, configured templates are saved permanently in the template library. Generated output files are still download-only and not retained.</p>
        """,
    )

    template_name = st.text_input("Template name", placeholder="Example: Urgent Care Pricing Agreement")
    category = st.text_input("Template category", placeholder="Example: United States - Urgent Care")
    requires_code = st.toggle("Requires automatic overseas currency code", value=False)
    description = st.text_area("Optional description", height=95)
    uploaded_template = st.file_uploader("Upload .docx template", type=["docx"], key="custom_template_upload")

    if uploaded_template is not None:
        template_bytes = uploaded_template.getvalue()
        try:
            counts = scan_template(template_bytes)
            st.markdown("### Template validation")
            metrics_card(counts)

            if counts["name"] == 0 or counts["address"] == 0 or counts["date"] == 0:
                st.warning("This template may be missing required yellow, green, or blue placeholders.")
            if requires_code and counts["code"] == 0:
                st.warning("This template is marked as requiring an automatic overseas currency code, but no gray placeholders were found.")

            neon_ready = database_configured()
            save_label = "Save template to Neon library" if neon_ready else "Save template for this session"
            if st.button(save_label, type="primary"):
                if not template_name.strip():
                    st.error("Enter a template name before saving.")
                elif not category.strip():
                    st.error("Enter a template category before saving.")
                elif neon_ready:
                    try:
                        save_template_to_neon(
                            name=template_name.strip(),
                            category=category.strip(),
                            requires_code=requires_code,
                            description=description.strip(),
                            filename=safe_filename(uploaded_template.name),
                            template_bytes=template_bytes,
                        )
                        st.success("Template saved permanently to Neon. It is now available in Generate Agreements.")
                        st.rerun()
                    except Exception as exc:
                        st.error(f"Could not save to Neon: {exc}")
                else:
                    custom = TemplateMeta(
                        id=f"session-{slugify(template_name)}-{uuid.uuid4().hex[:8]}",
                        name=template_name.strip(),
                        category=category.strip(),
                        requires_code=requires_code,
                        description=description.strip(),
                        source="session",
                        filename=safe_filename(uploaded_template.name),
                        bytes_data=template_bytes,
                        file_sha256=sha256_bytes(template_bytes),
                    )
                    st.session_state.setdefault("custom_templates", []).append(custom)
                    st.warning("Template saved for this browser session only because Neon is not connected.")
        except Exception as exc:
            st.error(f"Could not validate this template: {exc}")

    if templates:
        st.markdown("### Template library")
        library_df = pd.DataFrame(
            [
                {
                    "template_name": item.name,
                    "category": item.category,
                    "requires_currency_code": "Yes" if item.requires_code else "No",
                    "source": item.source,
                    "default": "Yes" if item.is_default else "No",
                    "description": item.description,
                }
                for item in templates
            ]
        )
        st.dataframe(library_df, use_container_width=True, hide_index=True)


def main() -> None:
    st.set_page_config(page_title=APP_TITLE, page_icon="📄", layout="wide")
    inject_css()

    st.markdown(
        """
        <div class="hero-card">
            <h1 class="hero-title">Pricing Agreement Generator</h1>
            <p class="hero-subtitle">Generate clean provider agreements from permanently saved Word templates. Upload a spreadsheet or paste providers, let the app resolve overseas currency codes, then download everything as a ZIP.</p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    generate_tab, configure_tab = st.tabs(["Generate Agreements", "Configure Templates"])
    with generate_tab:
        render_generate_tab()
    with configure_tab:
        render_configure_tab()


if __name__ == "__main__":
    main()
